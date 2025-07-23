import os
import json
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import exifread
import PyPDF2
import zipfile
from docx import Document
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
import hashlib
import datetime

console = Console()

class MetadataExtractor:
    def __init__(self):
        self.metadata = {}

    def extract(self, filepath):
        if not os.path.exists(filepath):
            console.print(f"[red]File not found: {filepath}[/red]")
            return

        console.print(f"[red]Extracting metadata from: {filepath}[/red]")
        
        file_ext = os.path.splitext(filepath)[1].lower()
        
        self.get_file_info(filepath)
        
        if file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.gif', '.bmp']:
            self.extract_image_metadata(filepath)
        elif file_ext == '.pdf':
            self.extract_pdf_metadata(filepath)
        elif file_ext in ['.docx', '.doc']:
            self.extract_docx_metadata(filepath)
        elif file_ext in ['.mp4', '.avi', '.mov', '.mkv']:
            self.extract_video_metadata(filepath)
        else:
            console.print(f"[yellow]Unsupported file type: {file_ext}[/yellow]")
        
        self.display_results()

    def get_file_info(self, filepath):
        stat = os.stat(filepath)
        
        self.metadata['File Info'] = {
            'Filename': os.path.basename(filepath),
            'File Size': f"{stat.st_size:,} bytes",
            'Created': datetime.datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
            'Modified': datetime.datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
            'Accessed': datetime.datetime.fromtimestamp(stat.st_atime).strftime('%Y-%m-%d %H:%M:%S'),
        }
        
        with open(filepath, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()
            self.metadata['File Info']['MD5 Hash'] = file_hash

    def extract_image_metadata(self, filepath):
        console.print("[cyan]Extracting image metadata...[/cyan]")
        
        try:
            with Image.open(filepath) as image:
                exif_data = image._getexif()
                
                if exif_data:
                    self.metadata['EXIF Data'] = {}
                    
                    for tag_id, value in exif_data.items():
                        tag = TAGS.get(tag_id, tag_id)
                        
                        if tag == 'GPSInfo':
                            gps_data = {}
                            for gps_tag_id, gps_value in value.items():
                                gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                                gps_data[gps_tag] = gps_value
                            
                            if gps_data:
                                coords = self.parse_gps_coordinates(gps_data)
                                if coords:
                                    self.metadata['GPS Location'] = coords
                        else:
                            self.metadata['EXIF Data'][tag] = str(value)[:100]
                
                self.metadata['Image Info'] = {
                    'Format': image.format,
                    'Mode': image.mode,
                    'Size': f"{image.size[0]}x{image.size[1]}",
                }
        
        except Exception as e:
            console.print(f"[red]Error reading image metadata: {e}[/red]")
        
        try:
            with open(filepath, 'rb') as f:
                tags = exifread.process_file(f, details=False)
                
                if tags:
                    self.metadata['Additional EXIF'] = {}
                    for tag, value in tags.items():
                        if not tag.startswith('JPEGThumbnail'):
                            self.metadata['Additional EXIF'][tag] = str(value)[:100]
        
        except Exception as e:
            pass

    def parse_gps_coordinates(self, gps_data):
        try:
            lat_ref = gps_data.get('GPSLatitudeRef')
            lat = gps_data.get('GPSLatitude')
            lon_ref = gps_data.get('GPSLongitudeRef')
            lon = gps_data.get('GPSLongitude')
            
            if lat and lon:
                lat_decimal = self.convert_to_decimal(lat, lat_ref)
                lon_decimal = self.convert_to_decimal(lon, lon_ref)
                
                return {
                    'Latitude': lat_decimal,
                    'Longitude': lon_decimal,
                    'Google Maps': f"https://maps.google.com/?q={lat_decimal},{lon_decimal}"
                }
        except:
            pass
        return None

    def convert_to_decimal(self, coord, ref):
        degrees = float(coord[0])
        minutes = float(coord[1])
        seconds = float(coord[2])
        
        decimal = degrees + minutes/60 + seconds/3600
        
        if ref in ['S', 'W']:
            decimal = -decimal
            
        return decimal

    def extract_pdf_metadata(self, filepath):
        console.print("[cyan]Extracting PDF metadata...[/cyan]")
        
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if pdf_reader.metadata:
                    self.metadata['PDF Metadata'] = {}
                    for key, value in pdf_reader.metadata.items():
                        clean_key = key.replace('/', '')
                        self.metadata['PDF Metadata'][clean_key] = str(value)[:100]
                
                self.metadata['PDF Info'] = {
                    'Pages': len(pdf_reader.pages),
                    'Encrypted': pdf_reader.is_encrypted,
                }
                
                if len(pdf_reader.pages) > 0:
                    first_page = pdf_reader.pages[0]
                    text_sample = first_page.extract_text()[:200]
                    if text_sample.strip():
                        self.metadata['Content Sample'] = {'First Page Text': text_sample}
        
        except Exception as e:
            console.print(f"[red]Error reading PDF metadata: {e}[/red]")

    def extract_docx_metadata(self, filepath):
        console.print("[cyan]Extracting DOCX metadata...[/cyan]")
        
        try:
            doc = Document(filepath)
            core_props = doc.core_properties
            
            self.metadata['Document Properties'] = {}
            
            properties = [
                'title', 'author', 'subject', 'keywords', 'comments', 
                'category', 'created', 'modified', 'last_modified_by',
                'revision', 'version'
            ]
            
            for prop in properties:
                value = getattr(core_props, prop, None)
                if value:
                    self.metadata['Document Properties'][prop.replace('_', ' ').title()] = str(value)
            
            if doc.paragraphs:
                text_sample = ""
                for para in doc.paragraphs[:5]:
                    text_sample += para.text + " "
                if text_sample.strip():
                    self.metadata['Content Sample'] = {'Text Sample': text_sample[:200]}
        
        except Exception as e:
            console.print(f"[red]Error reading DOCX metadata: {e}[/red]")

    def extract_video_metadata(self, filepath):
        console.print("[cyan]Video metadata extraction not implemented yet[/cyan]")

    def display_results(self):
        console.print(f"\n[red]═══ METADATA ANALYSIS REPORT ═══[/red]")
        
        for category, data in self.metadata.items():
            if data:
                table = Table(title=category, border_style="red")
                table.add_column("Property", style="cyan")
                table.add_column("Value", style="white")
                
                for key, value in data.items():
                    table.add_row(str(key), str(value))
                
                console.print(table)
                console.print()
        
        if 'GPS Location' in self.metadata:
            console.print("[yellow]⚠️  GPS coordinates found - potential privacy risk![/yellow]")
        
        if any('Author' in str(data) or 'Creator' in str(data) for data in self.metadata.values()):
            console.print("[yellow]⚠️  Author/Creator information found![/yellow]")
        
        console.print(f"[green]Metadata extraction complete[/green]")
